
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv
from openai import OpenAI
import tempfile
import pypandoc

# 🔐 Chargement clé API OpenAI
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# 🔧 Base dir robuste
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.."))

DATA_PATH = os.path.join(BASE_DIR, "data", "donnees_nettoyees.csv")
RESUME_PATH = os.path.join(BASE_DIR, "data", "resume_donnees.txt")
IMG_DIR = os.path.join(BASE_DIR, "data", "visuels")
OUTPUT_DIR = os.path.join(BASE_DIR, "outputs", "rapports")

# ✅ Images : tous les graphiques
GRAPH_NAMES = [
    "evolution_flux.png",
    "top_produits.png",
    "montant_par_annee.png",
    "heatmap_produits_annees.png",
    "evolution_top5.png"
]

def generer_resume_donnees(df):
    resume = df.groupby(["Année", "Flux"]).agg({"Montant": "sum"}).reset_index()
    resume_txt = resume.to_string(index=False)

    with open(RESUME_PATH, "w", encoding="utf-8") as f:
        f.write(resume_txt)

    return resume_txt

def generer_contenu_gpt(resume):
    prompt = f"""
    Tu es un analyste économique sénégalais.
    Rédige un rapport structuré et professionnel sur les échanges commerciaux du Sénégal.

    Structure :
    I. Introduction
    II. Analyse globale des flux
    III. Produits majeurs
    IV. Tendances principales
    V. Conclusion

    Voici le résumé des données :
    {resume}
    """
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Tu es un expert en analyse économique."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.7
    )
    return response.choices[0].message.content

def generer_rapport_docx(texte, images=[], titre="Rapport économique du Sénégal"):
    doc = Document()
    doc.add_heading(titre, level=1)

    for ligne in texte.split('\n'):
        if ligne.strip():
            doc.add_paragraph(ligne.strip())

    for image_path in images:
        if os.path.exists(image_path):
            doc.add_paragraph("\n")
            doc.add_picture(image_path, width=Inches(5.5))

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=OUTPUT_DIR)
    doc.save(temp_docx.name)
    return temp_docx.name

def convertir_en_pdf(docx_path):
    try:
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf", dir=OUTPUT_DIR)
        pypandoc.convert_file(docx_path, 'pdf', outputfile=temp_pdf.name)
        return temp_pdf.name
    except Exception as e:
        print("⚠️ Conversion PDF échouée :", e)
        return None

def generer_rapport_complet():
    df = pd.read_csv(DATA_PATH)
    resume = generer_resume_donnees(df)
    texte = generer_contenu_gpt(resume)

    images = [os.path.join(IMG_DIR, name) for name in GRAPH_NAMES]
    docx_path = generer_rapport_docx(texte, images=images)
    pdf_path = convertir_en_pdf(docx_path)
    return docx_path, pdf_path
